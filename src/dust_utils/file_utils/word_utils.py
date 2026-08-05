import re
from pathlib import Path

# 初始化日志
from loguru import logger


class WordUtils:
    """
    word工具类
    使用库：python-docx
    """

    @staticmethod
    def replace_vars(docx_path: str, params: dict):
        """
        替换word文档中的变量：支持 {{ var }} / {{ va r }} / {{Var}} 等模糊变量替换

        Args:
            docx_path (str): 输入 Word 文档的路径 (.docx)
            params (dict): 形如{{var}}，变量替换字典，键为变量名，值为替换值
        """
        logger.info(f"开始替换文档变量\n路径： {docx_path}\n变量：{params}")
        from docx import Document

        doc = Document(docx_path)

        # 处理普通段落
        for p in doc.paragraphs:
            WordUtils._process_paragraph(p, params)

        # 处理表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    WordUtils._process_cell(cell, params)

        doc.save(docx_path)
        logger.success(f"替换文档变量完成\n路径： {docx_path}")

    @staticmethod
    def _process_cell(cell: "_Cell", params: dict):
        for p in cell.paragraphs:
            WordUtils._process_paragraph(p, params)

    @staticmethod
    def _process_paragraph(p: "Paragraph", params: dict):
        """
        替换段落中的{{var}}为对应的值

        Args:
            p (Paragraph): 输入段落实例
            params (dict): 变量替换字典，键为变量名，值为替换值
        """
        if not p.runs:
            return

        # 拼接所有 run 的文本
        full_text = "".join(run.text for run in p.runs)

        # 判断是否包含变量
        if "{{" not in full_text:
            return

        # 变量替换
        full_text = WordUtils.replace_vars_fuzzy(full_text, params)

        # 取中间 run 的样式
        mid_index = len(p.runs) // 2
        src_run = p.runs[mid_index]

        # 清空原有 runs
        for run in p.runs:
            run.text = ""

        # 只保留一个 run
        new_run = p.add_run(full_text)

        # 复制样式
        new_run.bold = src_run.bold
        new_run.italic = src_run.italic
        new_run.underline = src_run.underline
        new_run.font.size = src_run.font.size
        new_run.font.color.rgb = src_run.font.color.rgb
        # 设置高亮
        highlight = WordUtils.get_highlight(src_run)
        new_run.font.highlight_color = highlight
        # 复制字体
        WordUtils._copy_font_name(src_run, new_run)

    @staticmethod
    def replace_vars_fuzzy(text: str, params: dict) -> str:
        """
        支持 {{ var }} / {{ va r }} / {{Var}} 等模糊变量替换

        Args:
            text (str): 输入文本实例
            params (dict): 变量替换字典，键为变量名，值为替换值
        """

        # 预处理参数：key 统一规范化
        norm_params = {re.sub(r"\s+", "", k).lower(): str(v) for k, v in params.items()}

        pattern = re.compile(r"\{\{(.*?)\}\}")

        def repl(match):
            raw_key = match.group(1)
            norm_key = re.sub(r"\s+", "", raw_key).lower()
            return norm_params.get(norm_key, match.group(0))

        return pattern.sub(repl, text)

    @staticmethod
    def _copy_font_name(src_run, target_run):
        """
        复制字体
        """
        target_run.font.name = src_run.font.name

        src_rPr = src_run._element.rPr
        if src_rPr is None:
            return

        rFonts = src_rPr.rFonts
        if rFonts is None:
            return

        from docx.oxml.ns import qn

        east_asia_font = rFonts.get(qn("w:eastAsia"))
        if not east_asia_font:
            return

        target_rPr = target_run._element.get_or_add_rPr()
        target_rFonts = target_rPr.get_or_add_rFonts()
        target_rFonts.set(qn("w:eastAsia"), east_asia_font)

    @staticmethod
    def get_highlight(run):
        """
        安全获取 run 的高亮颜色
        返回 WD_COLOR_INDEX 枚举或 None
        """
        try:
            from docx.enum.text import WD_COLOR_INDEX

            val = run.font.highlight_color
            # 如果不是合法枚举，返回 None
            if val in WD_COLOR_INDEX.__members__.values():
                return val
        except ValueError:
            # 遇到 'none' 或非法值
            return None
        return None

    @staticmethod
    def merge_docx(file_list, output_path):
        """
        按顺序合并多个 docx 文件

        :param file_list: list[str] docx文件路径
        :param output_file: 输出文件路径
        """

        # 懒加载第三方库
        from docx import Document
        from docxcompose.composer import Composer

        if not file_list:
            raise ValueError("合并文件列表不能为空")

        # 检测文件
        for file in file_list:
            path = Path(file)

            if not path.exists():
                raise FileNotFoundError(f"文件不存在: {file}")

            if path.suffix.lower() != ".docx":
                raise ValueError(f"不是docx文件: {file}")

        # 创建主文档
        master = Document(file_list[0])

        composer = Composer(master)

        # 依次追加
        for file in file_list[1:]:
            doc = Document(file)
            composer.append(doc)

        # 保存
        output_path = Path(output_path)

        output_path.parent.mkdir(parents=True, exist_ok=True)

        composer.save(str(output_path))
        return output_path
